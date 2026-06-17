"""附件解析：把 pdf / docx / md / txt 抽取为纯文本。

触发点：知识上线时由 Java 编排层带附件路径调 /ai/parse 或 /ai/index。
路径解析与 Java LocalFileStorage 对齐——相对路径在 settings.storage_dir 下解析。
"""
from __future__ import annotations

from html.parser import HTMLParser
from pathlib import Path

from app.core.config import settings
from app.core.logging import logger

# 支持的扩展名（小写，含点）
_SUPPORTED = {".txt", ".md", ".markdown", ".pdf", ".docx"}


def resolve_path(file_path: str) -> Path:
    """把 kb_attachment.file_path（相对 yyyy/MM/uuid.ext 或绝对）解析为绝对路径。"""
    p = Path(file_path)
    if not p.is_absolute():
        p = Path(settings.storage_dir) / p
    return p


def extract_text(file_path: str) -> str:
    """按扩展名分派抽取纯文本；不支持的类型抛 ValueError，文件不存在抛 FileNotFoundError。"""
    path = resolve_path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"附件不存在：{path}")

    ext = path.suffix.lower()
    if ext not in _SUPPORTED:
        raise ValueError(f"不支持的附件类型：{ext}（支持 {sorted(_SUPPORTED)}）")

    if ext == ".pdf":
        text = _from_pdf(path)
    elif ext == ".docx":
        text = _from_docx(path)
    elif ext in (".md", ".markdown"):
        text = _from_markdown(path)
    else:  # .txt
        text = _read_text(path)

    logger.info("解析附件 %s（%s）→ %d 字", path.name, ext, len(text))
    return text


def _read_text(path: Path) -> str:
    """读文本文件，优先 utf-8，回退 gbk（中文 Windows 常见），最后宽松忽略。"""
    raw = path.read_bytes()
    for enc in ("utf-8", "gbk"):
        try:
            return raw.decode(enc)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="ignore")


def _from_pdf(path: Path) -> str:
    from pypdf import PdfReader

    reader = PdfReader(str(path))
    parts = [(page.extract_text() or "") for page in reader.pages]
    return "\n".join(parts)


def _from_docx(path: Path) -> str:
    from docx import Document

    doc = Document(str(path))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    # 表格内容也纳入
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def _from_markdown(path: Path) -> str:
    """md → html → 去标签纯文本，剥离 markdown 标记噪声。"""
    import markdown as md

    html = md.markdown(_read_text(path))
    return _html_to_text(html)


class _TextExtractor(HTMLParser):
    def __init__(self) -> None:
        super().__init__()
        self._buf: list[str] = []

    def handle_data(self, data: str) -> None:
        self._buf.append(data)

    def text(self) -> str:
        return "".join(self._buf)


def _html_to_text(html: str) -> str:
    parser = _TextExtractor()
    parser.feed(html)
    return parser.text()
